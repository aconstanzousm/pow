#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para generar documento Word que explique el código del proyecto
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

# Crear documento
doc = Document()

# Configurar márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

# Página de titulo
title = doc.add_paragraph()
title_run = title.add_run('🔍 ANÁLISIS DETALLADO DEL CÓDIGO')
title_run.font.size = Pt(26)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 102, 204)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Backend y Frontend - Guía de Código Línea por Línea')
subtitle_run.font.size = Pt(14)
subtitle_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# PARTE 1: BACKEND
doc.add_heading('PARTE 1: BACKEND (Django REST API)', level=1)

# 1.1 Estructura General
doc.add_heading('1.1 Estructura General del Backend', level=2)
doc.add_paragraph('El backend sigue el patrón MTV (Model-Template-View) de Django:', style='List Bullet')

items = [
    'Models (modelos) - Define la estructura de datos en la BD',
    'Views (vistas) - Lógica de negocio y procesamiento',
    'URLs - Rutas que conectan requests con vistas',
    'Serializers - Convierten datos Python a JSON',
    'Settings - Configuración de la aplicación'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 1.2 Models.py
doc.add_heading('1.2 MODELS.PY - Definición de Datos', level=2)

doc.add_heading('¿Qué es un Modelo?', level=3)
doc.add_paragraph('Un modelo es una clase Python que representa una tabla en la base de datos. Cada atributo es una columna.')

doc.add_heading('Modelo: EMPLEADO', level=3)
code_empleado = '''class Empleado(models.Model):
    user = models.OneToOneField(...)           # Referencia única a usuario Django
    nombre = models.CharField(max_length=100)  # Texto hasta 100 caracteres
    apellido = models.CharField(max_length=100)
    email = models.EmailField(unique=True)     # Email único (no repetir)
    telefono = models.CharField(max_length=30, blank=True)  # Opcional
    activo = models.BooleanField(default=True)  # true/false
    created_at = models.DateTimeField(auto_now_add=True)   # Fecha de creación
    updated_at = models.DateTimeField(auto_now=True)       # Fecha actualización'''

p = doc.add_paragraph()
run = p.add_run(code_empleado)
run.font.family = 'Courier New'
run.font.size = Pt(8)

doc.add_paragraph()
doc.add_paragraph('Explicación:', style='List Bullet 2')
exp = [
    'OneToOneField: Una relación 1:1 con el usuario de Django (para login)',
    'CharField: Campo de texto, max_length limita caracteres',
    'EmailField: Valida que sea un email',
    'unique=True: No permite emails duplicados',
    'blank=True: Campo opcional',
    'auto_now_add: Se asigna automáticamente al crear',
    'auto_now: Se actualiza automáticamente cada vez que guarda'
]
for e in exp:
    doc.add_paragraph(e, style='List Bullet 3')

doc.add_paragraph()
doc.add_heading('Modelo: PRODUCTO', level=3)
code_producto = '''class Producto(models.Model):
    nombre = models.CharField(max_length=150)
    descripcion = models.CharField(max_length=255, blank=True)
    precio = models.PositiveIntegerField()     # Solo números positivos
    categoria = models.CharField(max_length=60, blank=True)
    activo = models.BooleanField(default=True)
    imagen = models.ImageField(upload_to="productos/", null=True, blank=True)
    stock = models.IntegerField(default=0)
    stock_minimo = models.PositiveIntegerField(default=5)  # Por defecto 5
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)'''

p = doc.add_paragraph()
run = p.add_run(code_producto)
run.font.family = 'Courier New'
run.font.size = Pt(8)

doc.add_paragraph()
doc.add_paragraph('Características importantes:', style='List Bullet 2')
chars = [
    'PositiveIntegerField: Solo acepta números ≥ 0',
    'ImageField: Almacena fotos (se guardan en media/productos/)',
    'default=0 y default=5: Valores iniciales',
    'stock_minimo: Usado para alertas de inventario bajo'
]
for c in chars:
    doc.add_paragraph(c, style='List Bullet 3')

doc.add_paragraph()
doc.add_heading('Modelo: PEDIDO', level=3)
code_pedido = '''class Pedido(models.Model):
    class Estado(models.TextChoices):
        PENDIENTE = "pendiente", "Pendiente"
        PREPARANDO = "preparando", "Preparando"
        LISTO = "listo", "Listo"
        CANCELADO = "cancelado", "Cancelado"

    cliente_nombre = models.CharField(max_length=150)
    cliente_email = models.EmailField(blank=True)
    nota = models.CharField(max_length=255, blank=True)
    estado = models.CharField(
        max_length=20,
        choices=Estado.choices,    # Solo acepta valores del enum
        default=Estado.PENDIENTE   # Comienza como PENDIENTE
    )
    total = models.PositiveIntegerField(default=0)
    empleado = models.ForeignKey(Empleado, null=True, blank=True, on_delete=models.SET_NULL)
    created_at = models.DateTimeField(auto_now_add=True)
    updated_at = models.DateTimeField(auto_now=True)'''

p = doc.add_paragraph()
run = p.add_run(code_pedido)
run.font.family = 'Courier New'
run.font.size = Pt(8)

doc.add_paragraph()
doc.add_paragraph('Concepto: Estados de Pedido', style='List Bullet 2')
doc.add_paragraph('TextChoices crea un enum limitado. Solo estos 4 valores son válidos.', style='List Bullet 3')
doc.add_paragraph('ForeignKey: Relación 1-a-muchos con Empleado (un empleado puede tener múltiples pedidos)', style='List Bullet 3')
doc.add_paragraph('on_delete=SET_NULL: Si se elimina empleado, el pedido queda con empleado=NULL', style='List Bullet 3')

doc.add_paragraph()
doc.add_heading('Modelo: PEDIDO_ITEM', level=3)
code_item = '''class PedidoItem(models.Model):
    pedido = models.ForeignKey(Pedido, related_name="items", on_delete=models.CASCADE)
    # related_name permite acceder así: pedido.items.all()
    
    producto = models.ForeignKey(Producto, null=True, blank=True, on_delete=models.SET_NULL)
    nombre = models.CharField(max_length=150)    # Snapshot del nombre
    precio = models.PositiveIntegerField()       # Snapshot del precio
    cantidad = models.PositiveIntegerField()     # Cuántas unidades'''

p = doc.add_paragraph()
run = p.add_run(code_item)
run.font.family = 'Courier New'
run.font.size = Pt(8)

doc.add_paragraph()
doc.add_paragraph('Propósito:', style='List Bullet 2')
doc.add_paragraph('PedidoItem es la tabla "intermedia" entre Pedido y Producto.', style='List Bullet 3')
doc.add_paragraph('Permite múltiples productos en un pedido. Ejemplo: 2 cafés + 1 medialuna.', style='List Bullet 3')
doc.add_paragraph('Guarda nombre y precio en ese momento (snapshot), no la referencia actual.', style='List Bullet 3')
doc.add_paragraph('on_delete=CASCADE: Si se elimina el pedido, se eliminan todos sus items.', style='List Bullet 3')

doc.add_page_break()

# 1.3 Views.py
doc.add_heading('1.3 VIEWS.PY - Lógica de Negocio', level=2)

doc.add_heading('Decoradores Importantes', level=3)

tabla_decoradores = doc.add_table(rows=5, cols=2)
tabla_decoradores.style = 'Light Grid Accent 1'

headers = tabla_decoradores.rows[0].cells
headers[0].text = 'Decorador'
headers[1].text = 'Qué Hace'

decs = [
    ['@api_view(["POST"])', 'Solo acepta requests POST'],
    ['@permission_classes([...])', 'Define quién puede acceder (autenticación)'],
    ['permissions.AllowAny()', 'Cualquiera puede acceder, sin login'],
    ['permissions.IsAdminUser()', 'Solo si es_admin (is_staff=True)']
]

for i, (dec, desc) in enumerate(decs, 1):
    row = tabla_decoradores.rows[i].cells
    row[0].text = dec
    row[1].text = desc

doc.add_paragraph()
doc.add_heading('Función: login_view', level=3)

code_login = '''@api_view(['POST'])
@permission_classes([permissions.AllowAny])
def login_view(request):
    # 1. Obtener datos del request (usuario y contraseña)
    username = request.data.get('username', '').strip()
    password = request.data.get('password', '').strip()
    
    # 2. Validar que no estén vacíos
    if not username or not password:
        return Response(
            {'success': False, 'error': 'Username and password are required'},
            status=status.HTTP_400_BAD_REQUEST  # Código de error 400
        )
    
    # 3. Usar Django Auth para autenticar
    user = authenticate(username=username, password=password)
    
    # 4. Si es None, credenciales incorrectas
    if user is None:
        return Response(
            {'success': False, 'error': 'Invalid username or password'},
            status=status.HTTP_401_UNAUTHORIZED  # Código 401
        )
    
    # 5. Determinar rol: es admin si is_staff=True
    role = 'admin' if user.is_staff else 'empleado'
    
    # 6. Retornar JSON con datos del usuario
    return Response({
        'success': True,
        'role': role,
        'username': user.username,
        'user_id': user.id
    }, status=status.HTTP_200_OK)  # Código 200 = éxito'''

p = doc.add_paragraph()
run = p.add_run(code_login)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_paragraph()
doc.add_heading('ViewSets: Operaciones CRUD Automáticas', level=3)

doc.add_paragraph('Un ViewSet es una clase que genera automáticamente las 5 operaciones:', style='List Bullet 2')

crud_ops = [
    'GET /productos/ → list() - Listar todos',
    'POST /productos/ → create() - Crear nuevo',
    'GET /productos/{id}/ → retrieve() - Obtener por ID',
    'PATCH /productos/{id}/ → update() - Modificar',
    'DELETE /productos/{id}/ → destroy() - Eliminar'
]
for op in crud_ops:
    doc.add_paragraph(op, style='List Bullet 3')

doc.add_paragraph()
code_viewset = '''class ProductoViewSet(viewsets.ModelViewSet):
    queryset = Producto.objects.all().order_by("-id")  # Ordenar por ID descendente
    serializer_class = ProductoSerializer
    
    def get_permissions(self):
        # Permiso diferenciado según la acción
        if self.action in ('list', 'retrieve'):
            return [permissions.AllowAny()]  # Ver menú = sin login
        return [permissions.IsAdminUser()]   # Crear/editar = solo admin
    
    def perform_create(self, serializer):
        # Se ejecuta al crear (POST)
        if not self.request.user.is_staff:
            raise PermissionDenied()  # Bloquear si no es admin
        serializer.save()'''

p = doc.add_paragraph()
run = p.add_run(code_viewset)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_page_break()

# 1.4 Serializers
doc.add_heading('1.4 SERIALIZERS.PY - Conversión de Datos', level=2)

doc.add_heading('¿Qué es un Serializer?', level=3)
doc.add_paragraph('Convierte objetos Python (modelos) a JSON y viceversa.', style='List Bullet 2')
doc.add_paragraph('Es como un "traductor" entre la base de datos y el frontend.', style='List Bullet 3')

doc.add_paragraph()
code_serializer = '''class ProductoSerializer(serializers.ModelSerializer):
    imagen = serializers.ImageField(
        required=False,          # Campo opcional
        allow_null=True,
        use_url=True            # Devuelve URL, no archivo
    )
    bajo_stock = serializers.SerializerMethodField(read_only=True)
    
    class Meta:
        model = Producto
        fields = (
            "id", "nombre", "descripcion", "precio", 
            "categoria", "activo", "imagen", 
            "stock", "stock_minimo", "bajo_stock",
            "created_at", "updated_at",
        )
    
    def get_bajo_stock(self, obj):
        # Calcula si stock <= stock_minimo (alerta)
        return obj.stock <= obj.stock_minimo'''

p = doc.add_paragraph()
run = p.add_run(code_serializer)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_paragraph()
doc.add_paragraph('Ejemplo JSON que devuelve:', style='List Bullet 2')

json_ejemplo = '''{
    "id": 1,
    "nombre": "Café Espresso",
    "descripcion": "Café fuerte y aromático",
    "precio": 3500,
    "categoria": "Bebidas",
    "activo": true,
    "imagen": "http://localhost:8000/media/productos/cafe.jpg",
    "stock": 15,
    "stock_minimo": 5,
    "bajo_stock": false,
    "created_at": "2026-06-18T10:30:00Z",
    "updated_at": "2026-06-18T14:20:00Z"
}'''

p = doc.add_paragraph()
run = p.add_run(json_ejemplo)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_page_break()

# 1.5 URLs
doc.add_heading('1.5 URLS.PY - Enrutamiento', level=2)

code_urls = '''from rest_framework.routers import DefaultRouter

router = DefaultRouter()
router.register(r"pedidos", PedidoViewSet, basename="pedido")
router.register(r"productos", ProductoViewSet, basename="producto")
router.register(r"empleados", EmpleadoViewSet, basename="empleado")

urlpatterns = [
    # Función: login
    path("login/", login_view, name="login"),
    
    # Vistas HTML (para acceso directo en navegador)
    path("", IndexView.as_view(), name="index"),
    path("menu/", MenuView.as_view(), name="menu"),
    path("checkout/", CheckoutView.as_view(), name="checkout"),
    
    # Rutas de la API (generadas automáticamente por router)
    path("", include(router.urls)),
    # Genera:
    # GET    /productos/        → list todos
    # POST   /productos/        → crear
    # GET    /productos/{id}/   → obtener
    # PATCH  /productos/{id}/   → editar
    # DELETE /productos/{id}/   → eliminar
]'''

p = doc.add_paragraph()
run = p.add_run(code_urls)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_paragraph()
doc.add_heading('DefaultRouter automáticamente crea:', level=3)
rutas = [
    'GET /productos/ - Lista todos (sin autenticación)',
    'POST /productos/ - Crear (solo admin)',
    'GET /productos/{id}/ - Obtener uno (sin autenticación)',
    'PATCH /productos/{id}/ - Editar (solo admin)',
    'DELETE /productos/{id}/ - Eliminar (solo admin)',
    '(Lo mismo para /pedidos/ y /empleados/)'
]
for ruta in rutas:
    doc.add_paragraph(ruta, style='List Bullet')

doc.add_page_break()

# PARTE 2: FRONTEND
doc.add_heading('PARTE 2: FRONTEND (HTML/CSS/JavaScript)', level=1)

doc.add_heading('2.1 Estructura General del Frontend', level=2)

doc.add_paragraph('El frontend está formado por:', style='List Bullet')
front_items = [
    'HTML (.html) - Estructura y contenido',
    'CSS (.css) - Estilos y diseño',
    'JavaScript - Interactividad y comunicación con API',
    'Assets - Imágenes y recursos'
]
for item in front_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Páginas HTML del Proyecto:', level=3)
paginas = [
    'index.html - Página de inicio con hero',
    'menu.html - Catálogo de productos',
    'checkout.html - Carrito y resumen de compra',
    'payment.html - Procesamiento de pago',
    'aboutus.html - Información sobre el café',
    'admin/admin-login.html - Login para administradores',
    'admin/admin-panel.html - Panel de control (admin)',
    'admin/empleado-panel.html - Panel de control (empleados)'
]
for pag in paginas:
    doc.add_paragraph(pag, style='List Bullet')

doc.add_page_break()

doc.add_heading('2.2 HTML - Estructura y Plantillas', level=2)

doc.add_heading('Ejemplo: menu.html (Estructura)', level=3)

code_html = '''<!DOCTYPE html>
<html lang="es">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Café Usm | Menú</title>
    <link rel="stylesheet" href="{% static 'menu.css' %}">
    <!-- {% static %} es un template tag de Django -->
    <!-- Busca el archivo en la carpeta static/ -->
</head>
<body>
    <!-- Navegación -->
    <nav class="glass-nav">
        <div class="logo">CAFÉ<span>USM</span></div>
        <ul class="nav-links">
            <li><a href="index.html">Inicio</a></li>
            <li><a href="menu.html" class="active">Menú</a></li>
        </ul>
    </nav>
    
    <!-- Container principal -->
    <main class="menu-container">
        <section class="menu-header">
            <h1>Nuestro Menú</h1>
        </section>
        
        <!-- Grid de productos (se llena con JavaScript) -->
        <div class="items-grid" id="productsGrid">
            <!-- Los productos se agregan aquí dinámicamente -->
        </div>
    </main>
    
    <!-- Carrito flotante -->
    <aside class="cart-sidebar" id="cartSidebar">
        <h3>🛒 Carrito</h3>
        <div id="cartItems"><!-- Items del carrito aquí --></div>
        <div id="cartTotal">Total: $0</div>
        <button onclick="goToCheckout()">Proceder al Pago</button>
    </aside>
    
    <script src="{% static 'menu.js' %}"></script>
</body>
</html>'''

p = doc.add_paragraph()
run = p.add_run(code_html)
run.font.family = 'Courier New'
run.font.size = Pt(6)

doc.add_paragraph()
doc.add_heading('Conceptos Clave del HTML:', level=3)
conceptos = [
    '{% load static %} - Activa el tag de Django para archivos estáticos',
    '{% static "archivo.css" %} - Obtiene la URL correcta del archivo',
    'id="productsGrid" - Identificador único para acceder en JavaScript',
    'class="items-grid" - Clase para aplicar estilos CSS',
    'onclick="function()" - Ejecuta función JavaScript al hacer click',
    'Document Object Model (DOM) - Tree de elementos HTML accesible desde JS'
]
for c in conceptos:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

doc.add_heading('2.3 CSS - Estilos y Variables', level=2)

doc.add_heading('Variables CSS (Temas)', level=3)

code_css = ''':root {
    /* Colores */
    --color-1: #624e88;      /* Púrpura principal */
    --color-2: #dc6ba0;      /* Rosa */
    --color-3: #cb80ab;      /* Rosa claro */
    --text-muted: #999;      /* Gris para texto secundario */
    
    /* Efectos Glass (cristal) */
    --glass: rgba(255, 255, 255, 0.9);
    --glass-border: rgba(255, 255, 255, 0.6);
    
    /* Tipografía */
    --font-main: "Inter", sans-serif;
    --font-size-base: 16px;
}'''

p = doc.add_paragraph()
run = p.add_run(code_css)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_paragraph()
doc.add_heading('Estructura CSS del Proyecto:', level=3)
estructura_css = [
    'index.css - Estilos de página de inicio',
    'menu.css - Estilos del menú, productos y carrito',
    'admin.css - Estilos del panel administrativo',
    'Cada archivo incluye: Variables, Grid/Flexbox, Animaciones, Responsive'
]
for e in estructura_css:
    doc.add_paragraph(e, style='List Bullet')

doc.add_heading('Sistema de Grid CSS:', level=3)

code_grid = '''.items-grid {
    display: grid;
    grid-template-columns: repeat(auto-fill, minmax(220px, 1fr));
    gap: 1.5rem;  /* Espacio entre items */
}

/* Explicación:
   - display: grid → Sistema de grid automático
   - repeat(auto-fill, minmax(220px, 1fr))
     • auto-fill → Crea columnas automáticamente
     • minmax(220px, 1fr) → Cada columna mide 220px como mínimo,
       y expande para llenar el espacio disponible
   - gap: 1.5rem → Espacio de 1.5rem entre elementos
   
   Resultado: En pantalla grande, 4-5 productos por fila
             En móvil, 1-2 productos por fila
*/'''

p = doc.add_paragraph()
run = p.add_run(code_grid)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_page_break()

doc.add_heading('2.4 JavaScript - Interactividad', level=2)

doc.add_heading('Flujo General de JavaScript', level=3)

flujo = [
    '1. Al cargar menu.html, se ejecuta el código en menu.js',
    '2. Se llama a fetchProductos() para obtener datos de la API',
    '3. La API retorna JSON con todos los productos',
    '4. Se recorren los productos con forEach()',
    '5. Se crea HTML dinámicamente para cada producto',
    '6. Se inserta el HTML en el DOM (document.getElementById("productsGrid"))',
    '7. Usuario hace click en "Agregar al carrito"',
    '8. Se llama addToCart(productId, cantidad)',
    '9. Se guarda en localStorage (navegador)',
    '10. Se actualiza la visualización del carrito'
]
for f in flujo:
    doc.add_paragraph(f, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Funciones JavaScript Principales:', level=3)

doc.add_paragraph('1. fetchProductos() - Obtener productos de la API', style='List Bullet 2')
code_fetch = '''async function fetchProductos() {
    try {
        const response = await fetch('/productos/');
        // fetch() hace request HTTP a /productos/
        
        if (!response.ok) {
            throw new Error('Error al obtener productos');
        }
        
        const productos = await response.json();
        // response.json() convierte JSON en objeto JavaScript
        
        renderProductos(productos);
        // Llama función para mostrar en pantalla
        
    } catch (error) {
        console.error('Error:', error);
    }
}'''

p = doc.add_paragraph()
run = p.add_run(code_fetch)
run.font.family = 'Courier New'
run.font.size = Pt(6)

doc.add_paragraph()
doc.add_paragraph('2. renderProductos(productos) - Mostrar en pantalla', style='List Bullet 2')

code_render = '''function renderProductos(productos) {
    const grid = document.getElementById('productsGrid');
    // Obtiene el elemento con id="productsGrid"
    
    grid.innerHTML = '';  // Limpia contenido anterior
    
    productos.forEach(producto => {
        // Crea HTML para cada producto
        const html = `
            <div class="item-card">
                <img src="${producto.imagen}" alt="${producto.nombre}" 
                     class="product-img">
                <div class="item-info">
                    <h3>${producto.nombre}</h3>
                    <p>${producto.descripcion}</p>
                    <div class="item-right">
                        <span class="price">$${producto.precio}</span>
                        <button onclick="addToCart(${producto.id}, 1)">
                            Agregar
                        </button>
                    </div>
                </div>
            </div>
        `;
        
        grid.innerHTML += html;  // Agrega HTML al grid
    });
}'''

p = doc.add_paragraph()
run = p.add_run(code_render)
run.font.family = 'Courier New'
run.font.size = Pt(6)

doc.add_paragraph()
doc.add_paragraph('3. addToCart(productId, cantidad) - Agregar al carrito', style='List Bullet 2')

code_cart = '''function addToCart(productId, cantidad) {
    // Obtener carrito actual del localStorage (navegador)
    let cart = JSON.parse(localStorage.getItem('cart') || '{}');
    
    if (cart[productId]) {
        // Si el producto ya está en el carrito, aumentar cantidad
        cart[productId] += cantidad;
    } else {
        // Si no está, agregarlo
        cart[productId] = cantidad;
    }
    
    // Guardar carrito actualizado en localStorage
    localStorage.setItem('cart', JSON.stringify(cart));
    
    updateCartDisplay();  // Actualizar visualización
    showNotification('Producto agregado al carrito');
}'''

p = doc.add_paragraph()
run = p.add_run(code_cart)
run.font.family = 'Courier New'
run.font.size = Pt(6)

doc.add_page_break()

doc.add_heading('2.5 Comunicación Frontend-Backend (AJAX)', level=2)

doc.add_heading('Ejemplo: Crear un Pedido desde Checkout', level=3)

code_pedido_js = '''async function createPedido() {
    // 1. Obtener datos del formulario
    const clientName = document.getElementById('clientName').value;
    const clientEmail = document.getElementById('clientEmail').value;
    const nota = document.getElementById('nota').value;
    
    // 2. Obtener carrito del localStorage
    const cart = JSON.parse(localStorage.getItem('cart') || '{}');
    
    // 3. Calcular total
    let total = 0;
    const items = [];
    for (const [productId, cantidad] of Object.entries(cart)) {
        // Obtener precio del producto
        const product = await getProducto(productId);
        const subtotal = product.precio * cantidad;
        total += subtotal;
        
        items.push({
            "producto_id": parseInt(productId),
            "cantidad": cantidad
        });
    }
    
    // 4. Construir payload JSON para enviar al backend
    const payload = {
        "cliente_nombre": clientName,
        "cliente_email": clientEmail,
        "nota": nota,
        "total": total,
        "items": items
    };
    
    // 5. Enviar POST a la API
    try {
        const response = await fetch('/pedidos/', {
            method: 'POST',
            headers: {
                'Content-Type': 'application/json',
                'X-CSRFToken': getCookie('csrftoken')  // Seguridad Django
            },
            body: JSON.stringify(payload)
        });
        
        if (!response.ok) {
            const errorData = await response.json();
            alert('Error: ' + errorData.error);
            return;
        }
        
        const pedido = await response.json();
        // Backend retorna datos del pedido creado
        
        // 6. Limpiar carrito y redirigir
        localStorage.removeItem('cart');
        alert('Pedido creado: #' + pedido.id);
        window.location.href = '/payment.html?pedido=' + pedido.id;
        
    } catch (error) {
        console.error('Error al crear pedido:', error);
    }
}'''

p = doc.add_paragraph()
run = p.add_run(code_pedido_js)
run.font.family = 'Courier New'
run.font.size = Pt(6)

doc.add_paragraph()
doc.add_heading('Explicación paso a paso:', level=3)

pasos_api = [
    'fetch("/pedidos/", {...}) → Hace request POST a la API',
    'method: "POST" → Indica que es crear (no GET)',
    'headers → Información de la request (tipo de contenido, seguridad)',
    'body: JSON.stringify(payload) → Convierte objeto JS a JSON',
    'response.json() → Backend retorna JSON con datos del pedido creado',
    'localStorage → Almacenamiento local en el navegador (persiste)',
    'window.location.href → Redirige a otra página'
]

for paso in pasos_api:
    doc.add_paragraph(paso, style='List Bullet')

doc.add_page_break()

# FLUJO COMPLETO
doc.add_heading('3. FLUJO COMPLETO: De Cliente a Base de Datos', level=1)

doc.add_heading('Ejemplo: Cliente compra un café', level=2)

tabla_flujo = doc.add_table(rows=13, cols=3)
tabla_flujo.style = 'Light Grid Accent 1'

h_row = tabla_flujo.rows[0].cells
h_row[0].text = 'Paso'
h_row[1].text = 'Acción'
h_row[2].text = 'Tecnología'

flujo_data = [
    ['1', 'Cliente entra a menu.html', 'Navegador carga HTML'],
    ['2', 'JavaScript ejecuta fetchProductos()', 'JavaScript/async'],
    ['3', 'Se hace fetch() a /productos/', 'Request HTTP GET'],
    ['4', 'Django recibe request en ProductoViewSet.list()', 'Django'],
    ['5', 'Se consulta la BD: Producto.objects.all()', 'ORM de Django'],
    ['6', 'Se serializa a JSON con ProductoSerializer', 'Serializers'],
    ['7', 'API retorna JSON con productos', 'Response HTTP 200'],
    ['8', 'JavaScript renderiza HTML dinámicamente', 'DOM manipulation'],
    ['9', 'Usuario ve menú con productos', 'Frontend'],
    ['10', 'Usuario hace click en "Agregar"', 'Evento click'],
    ['11', 'addToCart() guarda en localStorage', 'localStorage'],
    ['12', 'Usuario procede a checkout', 'Frontend']
]

for i, (paso, accion, tech) in enumerate(flujo_data, 1):
    row = tabla_flujo.rows[i].cells
    row[0].text = paso
    row[1].text = accion
    row[2].text = tech

doc.add_paragraph()
doc.add_heading('Cuando Usuario confirma compra:', level=3)

tabla_compra = doc.add_table(rows=11, cols=3)
tabla_compra.style = 'Light Grid Accent 1'

h_row2 = tabla_compra.rows[0].cells
h_row2[0].text = 'Paso'
h_row2[1].text = 'Acción'
h_row2[2].text = 'BD Result'

compra_data = [
    ['1', 'createPedido() recolecta datos', '—'],
    ['2', 'Construye JSON con items', '—'],
    ['3', 'fetch(POST) a /pedidos/', '—'],
    ['4', 'Django recibe en PedidoViewSet.create()', '—'],
    ['5', 'Valida: ¿Café abierto? (CafeConfig)', 'config.abierto = true'],
    ['6', 'Crea registro Pedido', 'INSERT INTO pedido ...'],
    ['7', 'Crea registros PedidoItem (para cada item)', 'INSERT INTO pedido_item ...'],
    ['8', 'Valida y ajusta stock', 'UPDATE producto SET stock = stock - cantidad'],
    ['9', 'Retorna JSON con pedido creado', 'SELECT * FROM pedido WHERE id=..'],
    ['10', 'Frontend limpia carrito', 'localStorage.clear()'],
]

for i, (paso, accion, result) in enumerate(compra_data, 1):
    row = tabla_compra.rows[i].cells
    row[0].text = paso
    row[1].text = accion
    row[2].text = result

doc.add_page_break()

# RESUMEN TÉCNICO
doc.add_heading('4. RESUMEN TÉCNICO Y ARQUITECTURA', level=1)

doc.add_heading('Arquitectura de Tres Capas:', level=2)

doc.add_paragraph('Frontend (Presentación)', style='List Bullet 2')
front = [
    'Capa: HTML/CSS/JavaScript',
    'Responsabilidad: Mostrar interfaz al usuario',
    'Tecnología: navegador (Chrome, Firefox, etc)',
    'Ubicación: archivos .html, .css en carpeta /frontend'
]
for f in front:
    doc.add_paragraph(f, style='List Bullet 3')

doc.add_paragraph('Backend (Lógica de Negocio)', style='List Bullet 2')
back = [
    'Capa: Django REST API',
    'Responsabilidad: Procesar requests, validar datos',
    'Tecnología: Python, Django, Django REST Framework',
    'Ubicación: views.py, models.py, serializers.py'
]
for b in back:
    doc.add_paragraph(b, style='List Bullet 3')

doc.add_paragraph('Base de Datos (Persistencia)', style='List Bullet 2')
db = [
    'Capa: SQLite/MySQL',
    'Responsabilidad: Almacenar datos de forma permanente',
    'Tecnología: SQL, Django ORM',
    'Ubicación: db.sqlite3'
]
for d in db:
    doc.add_paragraph(d, style='List Bullet 3')

doc.add_paragraph()
doc.add_heading('Patrón MTV (Model-Template-View) de Django:', level=2)

tabla_mtv = doc.add_table(rows=4, cols=3)
tabla_mtv.style = 'Light Grid Accent 1'

h_mtv = tabla_mtv.rows[0].cells
h_mtv[0].text = 'Componente'
h_mtv[1].text = 'Archivo'
h_mtv[2].text = 'Propósito'

mtv_data = [
    ['Model', 'models.py', 'Define estructura de BD'],
    ['View', 'views.py', 'Lógica de negocio'],
    ['Template', '*.html', 'Presentación (páginas)']
]

for i, (comp, archivo, prop) in enumerate(mtv_data, 1):
    row = tabla_mtv.rows[i].cells
    row[0].text = comp
    row[1].text = archivo
    row[2].text = prop

doc.add_page_break()

# CONCEPTOS FINALES
doc.add_heading('5. CONCEPTOS CLAVE A ENTENDER', level=1)

doc.add_heading('REST API - Principios Básicos:', level=2)

rest_conceptos = [
    'GET - Obtener datos (lectura)',
    'POST - Crear datos (escritura)',
    'PATCH - Modificar datos (actualización parcial)',
    'DELETE - Eliminar datos (borrado)',
    'Status Codes: 200 (OK), 201 (Creado), 400 (Error cliente), 401 (No autenticado), 403 (Prohibido), 404 (No existe), 500 (Error servidor)'
]
for c in rest_conceptos:
    doc.add_paragraph(c, style='List Bullet')

doc.add_heading('Async/Await en JavaScript:', level=2)

async_info = '''async function = función que retorna Promise
await = espera a que Promise se resuelva

Ejemplo:
    const response = await fetch('/productos/');
    // Espera hasta recibir respuesta de la API
    
    const data = await response.json();
    // Espera hasta convertir response a JSON
    
    console.log(data);  // Ahora data tiene los productos
'''

p = doc.add_paragraph()
run = p.add_run(async_info)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_heading('localStorage - Almacenamiento en Cliente:', level=2)

local_info = '''// Guardar
localStorage.setItem('cart', JSON.stringify(cart));

// Obtener
const cart = JSON.parse(localStorage.getItem('cart') || '{}');

// Limpiar
localStorage.removeItem('cart');

Usos: Carrito de compras, preferencias, tokens de sesión
Límite: ~5-10 MB por dominio
Persistencia: Se mantiene después de cerrar navegador
'''

p = doc.add_paragraph()
run = p.add_run(local_info)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_heading('ORM (Object-Relational Mapping):', level=2)

orm_info = '''Django ORM permite escribir SQL en Python:

En vez de SQL:
    SELECT * FROM productos WHERE stock > 10

Con Django ORM:
    Producto.objects.filter(stock__gt=10)

Ventajas:
- Código más legible
- Protección contra SQL injection
- Compatible con diferentes BDs
'''

p = doc.add_paragraph()
run = p.add_run(orm_info)
run.font.family = 'Courier New'
run.font.size = Pt(7)

doc.add_page_break()

# CONCLUSIÓN
doc.add_heading('6. CONCLUSIÓN', level=1)

conclusion = '''Este proyecto demuestra una arquitectura profesional fullstack:

BACKEND:
✓ Modelos bien definidos con relaciones
✓ ViewSets automáticos para CRUD
✓ Serializers para conversión de datos
✓ Permisos y autenticación
✓ Lógica de negocio centralizada

FRONTEND:
✓ HTML semántico
✓ CSS responsive con Grid
✓ JavaScript modular y async
✓ Comunicación AJAX con backend
✓ Almacenamiento local (localStorage)

FUNCIONALIDADES:
✓ Carrito persistente
✓ Estados de pedidos
✓ Control de stock
✓ Autenticación de roles
✓ Estadísticas en tiempo real

APRENDER ESTE PROYECTO TE DA HABILIDADES EN:
- Arquitectura fullstack
- API REST design
- Manejo de bases de datos
- Autenticación y autorización
- Frontend interactivo
- Comunicación cliente-servidor

¡Este código es listo para escala!
'''

doc.add_paragraph(conclusion)

doc.add_paragraph()
doc.add_paragraph('Última actualización: Junio 2026', style='List Bullet')
doc.add_paragraph('Versión: 1.0.0', style='List Bullet')

# Guardar
doc.save('CODIGO_EXPLICADO.docx')
print('✅ Documento "CODIGO_EXPLICADO.docx" creado correctamente')
