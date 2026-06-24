#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
Script para generar documentación en Word del proyecto Sistema de Gestión para Café
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Crear documento
doc = Document()

# Configurar márgenes
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Título principal
title = doc.add_paragraph()
title_run = title.add_run('📋 DOCUMENTACIÓN DEL PROYECTO')
title_run.font.size = Pt(28)
title_run.font.bold = True
title_run.font.color.rgb = RGBColor(0, 102, 204)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Subtítulo
subtitle = doc.add_paragraph()
subtitle_run = subtitle.add_run('Sistema Integral de Gestión para Café')
subtitle_run.font.size = Pt(16)
subtitle_run.font.italic = True
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()

# 1. DESCRIPCIÓN GENERAL
heading1 = doc.add_heading('1. DESCRIPCIÓN GENERAL', level=1)
heading1.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph('Este es un Sistema Integral de Gestión para Café que permite:', style='List Bullet')
features = [
    '👥 Gestionar productos (café, bebidas, alimentos)',
    '📦 Controlar inventario y stock',
    '📝 Registrar y seguir pedidos',
    '👤 Administrar empleados',
    '💰 Procesar pagos y generar reportes',
    '📊 Visualizar estadísticas del negocio'
]
for feature in features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Tecnología Principal:', level=3)
table = doc.add_table(rows=5, cols=2)
table.style = 'Light Grid Accent 1'

cells = table.rows[0].cells
cells[0].text = 'Componente'
cells[1].text = 'Tecnología'

data = [
    ['Backend', 'Django 5.2 + Django REST Framework'],
    ['Frontend', 'HTML5, CSS3, JavaScript vanilla'],
    ['Base de Datos', 'SQLite (desarrollo)'],
    ['Autenticación', 'Django Auth (usuario/contraseña)']
]

for i, (comp, tech) in enumerate(data, 1):
    row_cells = table.rows[i].cells
    row_cells[0].text = comp
    row_cells[1].text = tech

doc.add_page_break()

# 2. ESTRUCTURA DEL PROYECTO
heading2 = doc.add_heading('2. ESTRUCTURA DEL PROYECTO', level=1)
heading2.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph('Backend (API REST y lógica del servidor)', style='List Bullet 2')
backend_items = [
    'config/ - Configuraciones principales de Django',
    'coreapp/ - Aplicación principal (modelos, vistas, rutas)',
    'manage.py - Herramienta de gestión Django',
    'requirements.txt - Dependencias Python',
    'db.sqlite3 - Base de datos SQLite',
    'media/ - Imágenes de productos'
]
for item in backend_items:
    doc.add_paragraph(item, style='List Bullet 3')

doc.add_paragraph('Frontend (Interfaz de usuario)', style='List Bullet 2')
frontend_items = [
    'index.html - Página principal',
    'menu.html - Menú de productos',
    'checkout.html - Carrito de compras',
    'payment.html - Procesamiento de pagos',
    'aboutus.html - Acerca de nosotros',
    'admin/ - Panel administrativo (login, panel admin, panel empleado)',
]
for item in frontend_items:
    doc.add_paragraph(item, style='List Bullet 3')

doc.add_page_break()

# 3. BACKEND - MODELOS DE DATOS
heading3 = doc.add_heading('3. BACKEND - MODELOS DE DATOS', level=1)
heading3.runs[0].font.color.rgb = RGBColor(0, 102, 204)

# Empleado
doc.add_heading('3.1 Modelo: Empleado', level=2)
doc.add_paragraph('Representa a los trabajadores del café.')
empleado_attrs = [
    'nombre: Nombre del empleado',
    'apellido: Apellido',
    'email: Email único',
    'telefono: Número de contacto (opcional)',
    'activo: Si trabaja actualmente (true/false)',
    'user: Vinculación con usuario de autenticación',
    'created_at y updated_at: Timestamps'
]
for attr in empleado_attrs:
    doc.add_paragraph(attr, style='List Bullet')
doc.add_paragraph('Uso: Gestionar personal, autenticación, y atribución de pedidos a empleados.', style='List Bullet')
doc.add_paragraph()

# Producto
doc.add_heading('3.2 Modelo: Producto', level=2)
doc.add_paragraph('Artículos disponibles en el menú.')
producto_attrs = [
    'nombre: Nombre del producto (café, jugo, etc.)',
    'descripcion: Breve descripción',
    'precio: Costo en pesos (entero)',
    'categoria: Categoría (café, bebida, alimento)',
    'activo: Si está disponible para venta',
    'imagen: Foto del producto',
    'stock: Cantidad disponible',
    'stock_minimo: Cantidad mínima antes de alerta',
    'created_at y updated_at: Timestamps'
]
for attr in producto_attrs:
    doc.add_paragraph(attr, style='List Bullet')
doc.add_paragraph('Uso: Catálogo de productos que los clientes pueden comprar.', style='List Bullet')
doc.add_paragraph()

# Pedido
doc.add_heading('3.3 Modelo: Pedido', level=2)
doc.add_paragraph('Registro de compras de clientes.')
pedido_attrs = [
    'cliente_nombre: Nombre de quién compra',
    'cliente_email: Email del cliente (opcional)',
    'nota: Notas especiales o comentarios',
    'estado: PENDIENTE → PREPARANDO → LISTO → CANCELADO',
    'total: Monto total en pesos',
    'empleado: Quién prepara el pedido (opcional)',
    'created_at y updated_at: Timestamps'
]
for attr in pedido_attrs:
    doc.add_paragraph(attr, style='List Bullet')

doc.add_heading('Estados del Pedido:', level=3)
tabla_estados = doc.add_table(rows=5, cols=2)
tabla_estados.style = 'Light Grid Accent 1'

headers = tabla_estados.rows[0].cells
headers[0].text = 'Estado'
headers[1].text = 'Descripción'

estados = [
    ['PENDIENTE', 'Acababa de crearse, esperando preparación'],
    ['PREPARANDO', 'Alguien está haciendo el pedido'],
    ['LISTO', 'Completado, esperando retiro'],
    ['CANCELADO', 'Cancelado por el cliente']
]

for i, (estado, desc) in enumerate(estados, 1):
    row = tabla_estados.rows[i].cells
    row[0].text = estado
    row[1].text = desc

doc.add_paragraph()

# PedidoItem
doc.add_heading('3.4 Modelo: PedidoItem', level=2)
doc.add_paragraph('Productos individuales dentro de un pedido.')
item_attrs = [
    'pedido: Referencia al pedido principal',
    'producto: Referencia al producto',
    'nombre: Nombre del producto (copia en caso de cambios)',
    'precio: Precio en ese momento',
    'cantidad: Cuántas unidades'
]
for attr in item_attrs:
    doc.add_paragraph(attr, style='List Bullet')
doc.add_paragraph('Ejemplo: 1 café + 2 medialunas en un mismo pedido.', style='List Bullet')
doc.add_paragraph()

# CafeConfig
doc.add_heading('3.5 Modelo: CafeConfig', level=2)
doc.add_paragraph('Configuración global del café (singleton - solo 1 registro).')
config_attrs = [
    'abierto: Si el café está abierto o cerrado',
    'updated_at: Última vez que cambió'
]
for attr in config_attrs:
    doc.add_paragraph(attr, style='List Bullet')
doc.add_paragraph('Uso especial: Permite abrir/cerrar el café globalmente desde el panel admin.', style='List Bullet')

doc.add_page_break()

# 4. API REST
heading4 = doc.add_heading('4. API REST - ENDPOINTS PRINCIPALES', level=1)
heading4.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph('Base URL: http://localhost:8000/', style='List Bullet')

# Tabla de endpoints
doc.add_heading('Endpoints de Autenticación:', level=2)
tabla_auth = doc.add_table(rows=4, cols=3)
tabla_auth.style = 'Light Grid Accent 1'

h_row = tabla_auth.rows[0].cells
h_row[0].text = 'Método'
h_row[1].text = 'Endpoint'
h_row[2].text = 'Descripción'

auth_data = [
    ['POST', '/login/', 'Autenticar usuario'],
    ['GET/PATCH', '/cafe-status/', 'Ver/cambiar estado café'],
    ['GET', '/estadisticas/', 'Ver estadísticas (solo admin)']
]

for i, (metodo, endpoint, desc) in enumerate(auth_data, 1):
    row = tabla_auth.rows[i].cells
    row[0].text = metodo
    row[1].text = endpoint
    row[2].text = desc

doc.add_paragraph()
doc.add_heading('Endpoints CRUD (Completos):', level=2)

crud_items = [
    'GET /productos/ - Listar todos los productos',
    'POST /productos/ - Crear nuevo producto',
    'GET /productos/{id}/ - Ver detalles de producto',
    'PATCH /productos/{id}/ - Editar producto',
    'DELETE /productos/{id}/ - Eliminar producto'
]

for item in crud_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Los mismos endpoints existen para: /pedidos/ y /empleados/', style='List Bullet')

doc.add_page_break()

# 5. FRONTEND
heading5 = doc.add_heading('5. FRONTEND - PÁGINAS PRINCIPALES', level=1)
heading5.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_heading('5.1 Flujo de Usuario Cliente', level=2)
client_flow = [
    '1. Entra a index.html (página principal)',
    '2. Navega a menu.html (ve todos los productos)',
    '3. Selecciona productos → se agregan al carrito',
    '4. Va a checkout.html (revisa el carrito)',
    '5. Ingresa nombre y email',
    '6. Procede a payment.html (paga)',
    '7. ✅ Pedido creado (estado: PENDIENTE)',
    '8. Empleado lo ve y cambia estado'
]
for step in client_flow:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('5.2 Flujo de Usuario Empleado', level=2)
employee_flow = [
    '1. Accede a admin/empleado-panel.html',
    '2. Se autentica (login con usuario/contraseña)',
    '3. Ve pedidos en estado PENDIENTE',
    '4. Cambia a PREPARANDO (comienza a hacer)',
    '5. Cambia a LISTO (termina)',
    '6. Cliente retira → Pedido completado'
]
for step in employee_flow:
    doc.add_paragraph(step, style='List Bullet')

doc.add_heading('5.3 Flujo de Usuario Admin', level=2)
admin_flow = [
    '1. Accede a admin/admin-login.html',
    '2. Se autentica (solo si es admin: is_staff=True)',
    '3. En admin-panel.html puede:',
    '   • Ver estadísticas y gráficos de ventas',
    '   • Agregar/editar/eliminar productos',
    '   • Crear/eliminar empleados',
    '   • Ver todos los pedidos',
    '   • Abrir/cerrar el café',
    '   • Ver alertas de bajo stock'
]
for step in admin_flow:
    doc.add_paragraph(step, style='List Bullet')

doc.add_page_break()

# 6. BASE DE DATOS
heading6 = doc.add_heading('6. BASE DE DATOS - ESTRUCTURA SQLite', level=1)
heading6.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_paragraph('El proyecto usa SQLite (db.sqlite3) con las siguientes tablas:')

doc.add_heading('Tabla: EMPLEADO', level=3)
emp_fields = [
    'id (Clave Primaria)',
    'nombre (Texto)',
    'apellido (Texto)',
    'email (Texto - ÚNICO)',
    'telefono (Texto)',
    'activo (Booleano)',
    'user_id (Referencia a Django User)',
    'created_at y updated_at (Timestamps)'
]
for field in emp_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('Tabla: PRODUCTO', level=3)
prod_fields = [
    'id (Clave Primaria)',
    'nombre (Texto)',
    'descripcion (Texto)',
    'precio (Entero)',
    'categoria (Texto)',
    'activo (Booleano)',
    'imagen (Archivo)',
    'stock (Entero)',
    'stock_minimo (Entero)',
    'created_at y updated_at (Timestamps)'
]
for field in prod_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('Tabla: PEDIDO', level=3)
ped_fields = [
    'id (Clave Primaria)',
    'cliente_nombre (Texto)',
    'cliente_email (Texto)',
    'nota (Texto)',
    'estado (Enum: pendiente/preparando/listo/cancelado)',
    'total (Entero)',
    'empleado_id (Referencia a Empleado)',
    'created_at y updated_at (Timestamps)'
]
for field in ped_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_heading('Tabla: PEDIDO_ITEM', level=3)
item_fields = [
    'id (Clave Primaria)',
    'pedido_id (Referencia a Pedido)',
    'producto_id (Referencia a Producto)',
    'nombre (Texto - snapshot del producto)',
    'precio (Entero - snapshot del precio)',
    'cantidad (Entero)'
]
for field in item_fields:
    doc.add_paragraph(field, style='List Bullet')

doc.add_paragraph()

doc.add_heading('Migraciones de Base de Datos:', level=3)
migrations = [
    '0001_initial.py - Creación inicial de tablas',
    '0002_producto_imagen.py - Agregó campo imagen a Producto',
    '0003_cafeconfig.py - Agregó tabla de configuración',
    '0004_producto_stock.py - Agregó campos stock',
    '0005_producto_stock_minimo_alter_producto_stock.py - Stock mínimo'
]
for mig in migrations:
    doc.add_paragraph(mig, style='List Bullet')

doc.add_page_break()

# 7. INSTALACIÓN Y CONFIGURACIÓN
heading7 = doc.add_heading('7. GUÍA DE INSTALACIÓN Y CONFIGURACIÓN', level=1)
heading7.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_heading('Requisitos:', level=2)
reqs = [
    'Python 3.10 o superior',
    'pip (gestor de paquetes)',
    'Navegador web moderno'
]
for req in reqs:
    doc.add_paragraph(req, style='List Bullet')

doc.add_heading('Pasos de Instalación:', level=2)

doc.add_heading('1. Crear entorno virtual:', level=3)
p = doc.add_paragraph()
run = p.add_run('python -m venv wenv')
run.font.family = 'Courier New'
run.bold = True

doc.add_heading('2. Activar entorno virtual:', level=3)
p = doc.add_paragraph()
run = p.add_run('wenv\\Scripts\\activate')
run.font.family = 'Courier New'
run.bold = True
doc.add_paragraph('(en Windows)', style='List Bullet')

doc.add_heading('3. Instalar dependencias:', level=3)
p = doc.add_paragraph()
run = p.add_run('pip install -r backend/requirements.txt')
run.font.family = 'Courier New'
run.bold = True

doc.add_heading('4. Aplicar migraciones:', level=3)
p = doc.add_paragraph()
run = p.add_run('cd backend && python manage.py migrate')
run.font.family = 'Courier New'
run.bold = True

doc.add_heading('5. Crear superusuario (admin):', level=3)
p = doc.add_paragraph()
run = p.add_run('python manage.py createsuperuser')
run.font.family = 'Courier New'
run.bold = True

doc.add_heading('6. Iniciar servidor:', level=3)
p = doc.add_paragraph()
run = p.add_run('python manage.py runserver')
run.font.family = 'Courier New'
run.bold = True

doc.add_heading('7. Acceder a la aplicación:', level=3)
doc.add_paragraph('http://localhost:8000/', style='List Bullet')
doc.add_paragraph('Admin Django: http://localhost:8000/admin/', style='List Bullet')

doc.add_page_break()

# 8. SEGURIDAD
heading8 = doc.add_heading('8. SEGURIDAD', level=1)
heading8.runs[0].font.color.rgb = RGBColor(0, 102, 204)

sec_features = [
    '✅ Contraseñas encriptadas con Django (PBKDF2)',
    '✅ Autenticación por usuario y contraseña',
    '✅ Roles de usuario: Admin (is_staff) y Empleado',
    '✅ CORS (Cross-Origin Resource Sharing) configurado',
    '✅ Endpoints protegidos con permisos específicos',
    '✅ Solo admin puede ver estadísticas',
    '✅ Solo admin puede cambiar estado del café'
]
for feature in sec_features:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_page_break()

# 9. CASO DE USO COMPLETO
heading9 = doc.add_heading('9. CASO DE USO COMPLETO - EJEMPLO PRÁCTICO', level=1)
heading9.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_heading('Escenario: Cliente compra 2 cafés y 1 medialuna', level=2)

tabla_caso = doc.add_table(rows=11, cols=2)
tabla_caso.style = 'Light Grid Accent 1'

header_row = tabla_caso.rows[0].cells
header_row[0].text = 'Paso'
header_row[1].text = 'Acción'

caso_steps = [
    ['1', 'Cliente entra a index.html (página de inicio)'],
    ['2', 'Navega a menu.html y ve productos disponibles'],
    ['3', 'Café ($3.500), Medialuna ($2.000), Jugo ($2.500)'],
    ['4', 'Selecciona: 2 cafés + 1 medialuna → carrito'],
    ['5', 'Total en carrito: $9.000'],
    ['6', 'Procede a checkout.html'],
    ['7', 'Ingresa: nombre "Carlos", email "carlos@email.com"'],
    ['8', 'Va a payment.html y confirma pago'],
    ['9', 'Sistema crea PEDIDO #42 (estado: PENDIENTE)'],
    ['10', 'Empleado ve pedido en panel y comienza a preparar']
]

for i, (step, action) in enumerate(caso_steps, 1):
    row = tabla_caso.rows[i].cells
    row[0].text = step
    row[1].text = action

doc.add_heading('Continuación del Empleado:', level=3)
emp_continuation = [
    'Empleado en empleado-panel.html ve PEDIDO #42',
    'Cambia estado: PENDIENTE → PREPARANDO',
    'Prepara: 2 cafés + 1 medialuna',
    'Cambia estado: PREPARANDO → LISTO',
    'Notifica al cliente que está listo',
    'Cliente retira su pedido',
    '✅ PEDIDO COMPLETADO'
]
for item in emp_continuation:
    doc.add_paragraph(item, style='List Bullet')

doc.add_page_break()

# 10. PUNTOS CLAVE PARA LA PRESENTACIÓN
heading10 = doc.add_heading('10. PUNTOS CLAVE PARA LA PRESENTACIÓN', level=1)
heading10.runs[0].font.color.rgb = RGBColor(0, 102, 204)

doc.add_heading('✅ Aspectos Técnicos Importantes:', level=2)
tech_points = [
    'Arquitectura: Frontend - Backend - Base de Datos (separación de capas)',
    'API REST: Comunicación mediante HTTP (GET, POST, PATCH, DELETE)',
    'Base de Datos Relacional: Uso de claves foráneas y relaciones',
    'Autenticación: Sistema de roles (admin vs empleado)',
    'Estados del Pedido: Máquina de estados (PENDIENTE → LISTO → COMPLETADO)',
    'Stock Mínimo: Alertas automáticas para inventario bajo',
    'Timestamps: Auditoría de creación y modificación de datos'
]
for point in tech_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('✅ Funcionalidades Destacadas:', level=2)
features_destacadas = [
    'Carrito de compras persistente en frontend',
    'Panel administrativo con estadísticas en tiempo real',
    'Gestión de empleados y asignación de pedidos',
    'Control de inventario con alertas de bajo stock',
    'Sistema de estados de pedidos en tiempo real',
    'API REST completamente funcional y documentada'
]
for feature in features_destacadas:
    doc.add_paragraph(feature, style='List Bullet')

doc.add_heading('✅ Flujos de Negocio:', level=2)
business_flows = [
    'Cliente: Navega menú → Agrega productos → Compra',
    'Empleado: Ve pedidos → Prepara → Marca como listo',
    'Admin: Gestiona productos/empleados → Ve estadísticas',
    'Sistema: Actualiza stock → Genera alertas → Registra ingresos'
]
for flow in business_flows:
    doc.add_paragraph(flow, style='List Bullet')

doc.add_page_break()

# 11. PRÓXIMAS MEJORAS
heading11 = doc.add_heading('11. POSIBLES MEJORAS Y EXTENSIONES', level=1)
heading11.runs[0].font.color.rgb = RGBColor(0, 102, 204)

improvements = [
    'Integración de pagos real (Stripe, PayPal)',
    'Sistema de reservas para mesa',
    'Notificaciones en tiempo real (WebSocket)',
    'App móvil (React Native o Flutter)',
    'Sistema de reportes y exportación a Excel',
    'Integración con impresoras para recibos',
    'Sistema de promociones y descuentos',
    'Panel de análisis avanzado (BI)'
]
for improvement in improvements:
    doc.add_paragraph(improvement, style='List Bullet')

doc.add_page_break()

# Página final - Resumen
doc.add_heading('RESUMEN EJECUTIVO', level=1)

summary_text = '''Este proyecto es un sistema completo de gestión para un café que demuestra:

✅ Conocimientos de Desarrollo Full-Stack
• Backend con Django REST Framework
• Frontend interactivo con JavaScript vanilla
• Base de datos relacional con SQLite

✅ Conceptos de Ingeniería de Software
• Arquitectura cliente-servidor
• Separación de responsabilidades (MVC/MTV)
• API REST RESTful
• Gestión de estados y flujos de negocio

✅ Habilidades Prácticas
• Autenticación y autorización
• Gestión de base de datos relacional
• Validación de datos
• Manejo de errores
• Escalabilidad y mantenibilidad

✅ Aplicación Real
• Caso de uso práctico y realista
• Interfaz intuitiva para usuarios
• Lógica de negocio coherente
• Control de inventario funcional

El proyecto está listo para la producción con solo agregar:
• Integración de pagos real
• Hosting en servidor en la nube
• Certificado SSL/HTTPS
• Sistema de respaldos automáticos
'''

doc.add_paragraph(summary_text)

doc.add_paragraph()
doc.add_paragraph('Última actualización: Junio 2026', style='List Bullet')
doc.add_paragraph('Versión: 1.0.0', style='List Bullet')

# Guardar documento
doc.save('DOCUMENTACION_PROYECTO.docx')
print('✅ Archivo Word creado exitosamente: DOCUMENTACION_PROYECTO.docx')
